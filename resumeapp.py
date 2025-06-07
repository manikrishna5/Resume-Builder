import docx
from docx import Document
import streamlit as st
import os




def fill_resume_template(template_path, output_path, context):
    doc = Document(template_path)

    # Replace placeholders in paragraphs
    for para in doc.paragraphs:
        full_text = ''.join(run.text for run in para.runs)
        for key, val in context.items():
            if f"{{{{{key}}}}}" in full_text:
                full_text = full_text.replace(f"{{{{{key}}}}}", val)
        # Reassign full text back to runs
        if full_text != ''.join(run.text for run in para.runs):
            para.clear()  # Clear existing runs
            para.add_run(full_text)

    # Replace in tables if needed
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text = cell.text
                for key, val in context.items():
                    if f"{{{{{key}}}}}" in full_text:
                        full_text = full_text.replace(f"{{{{{key}}}}}", val)
                cell.text = full_text

    doc.save(output_path)


def generate_resume():
     if "resume_ready" not in st.session_state:
        st.session_state.resume_ready = False

     if "name" in st.session_state and st.session_state.get("shownext3"):
        st.subheader("✅ Final Step: Generate Your Resume")

        if st.button("📄 Generate Resume"):
            context = {
    "name": st.session_state.get("name", ""),
    "phone": st.session_state.get("phone", ""),
    "email": st.session_state.get("email", ""),
    "career": st.session_state.get("summary", ""),
    "college": st.session_state.get("college", ""),
    "cgp": st.session_state.get("cgpa", ""),
    "edu": st.session_state.get("education", ""),
    "branch": st.session_state.get("branch", ""),
    "startyear": str(st.session_state.get("startyear", "")),
    "endyear": str(st.session_state.get("endyear", "")),
    "languages_p": ', '.join(st.session_state.categorized_skills.get("Programming Languages", [])),
    "developer_tools": ', '.join(st.session_state.categorized_skills.get("Tools & Platforms", [])),
    "tech_frmae": ', '.join(st.session_state.categorized_skills.get("Frameworks / Libraries", [])),

    # 👇 NEW DYNAMIC FIELDS
    "all_experiences": "\n\n".join([
        f"{exp['company']} ({exp['start_date'].year}–{exp['end_date'].year})\n{exp['description']}"
        for exp in st.session_state.experiences
    ]) if st.session_state.experiences else "",

    "all_projects": "\n\n".join([
        f"{proj['title']} ({proj['link']})\n{proj['description']}"
        for proj in st.session_state.projects
    ]) if st.session_state.projects else "",

    "all_certifications": '\n'.join(st.session_state.get("certifications", [])),
    "descvol": st.session_state.get("volunteer_experience", ""),
    "all_interests": ', '.join(st.session_state.get("interests", [])),
    "all_languages": ', '.join(st.session_state.get("languages_known", [])),
}


            template_path = "sample_template_dynamic.docx"
            output_path = "generated_resume.docx"

            if not os.path.exists(template_path):
                st.error("❌ Template file not found. Please upload or place 'sample_template.docx' in the project directory.")
            else:
                fill_resume_template(template_path, output_path, context)
                st.session_state.resume_ready = True
                st.session_state.generated_resume_path = output_path
                st.session_state.generated_resume_name = f"{context['name'].replace(' ', '_')}_Resume.docx"
                st.success("🎉 Resume generated successfully! Scroll down to download.")

        if st.session_state.resume_ready and st.session_state.generated_resume_path:
            try:
                with open(st.session_state.generated_resume_path, "rb") as file:
                    st.download_button(
                        label="⬇️ Download Resume",
                        data=file,
                        file_name=st.session_state.generated_resume_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            except FileNotFoundError:
                st.error("⚠️ Resume file not found. Please try generating it again.")

